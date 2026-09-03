import io
import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from scripts import import_sharepoint_diagnostics as importer


ROOT = Path(__file__).resolve().parents[1]
SOURCE_NAME = "ХИМ_ОГЭ_Диагностика_21-22_Заданий 8.docx"


def _png(width: int = 40, height: int = 30) -> io.BytesIO:
    buffer = io.BytesIO()
    Image.new("RGB", (width, height), (10, 120, 200)).save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _task(document, number: int) -> None:
    document.add_paragraph(f"Задание {number}")


def _answer(document, solution: str | None, key: str) -> None:
    if solution is not None:
        document.add_paragraph("Решение:")
        document.add_paragraph(solution)
    document.add_paragraph("Ответ:")
    document.add_paragraph(key)


def build_source_document(path: Path) -> None:
    """One task per supported mapping, plus an irregular key and a figure."""
    document = Document()

    _task(document, 1)
    document.add_paragraph("Выберите одно вещество.")
    document.add_paragraph("Варианты:")
    for label in ("Кислород", "Азот", "Хлор"):
        document.add_paragraph(label)
    _answer(document, "Хлор стоит третьим в списке.", "3")

    _task(document, 2)
    document.add_paragraph("Выберите два вещества.")
    document.add_paragraph("Варианты:")
    for label in ("Медь", "Сера", "Железо", "Неон"):
        document.add_paragraph(label)
    _answer(document, None, "1#3")

    _task(document, 3)
    document.add_paragraph("Установите соответствие между формулой и классом.")
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "ФОРМУЛА"
    table.cell(0, 1).text = "КЛАСС"
    table.cell(1, 0).text = "А) HCl"
    table.cell(1, 1).text = "1) Кислота"
    table.cell(2, 0).text = "Б) NaOH"
    table.cell(2, 1).text = "2) Основание"
    _answer(document, "Соляная кислота и щёлочь.", "12")

    _task(document, 4)
    document.add_paragraph("Вычислите массовую долю в процентах.")
    _answer(document, "Считаем по формуле.", "0,25")

    _task(document, 5)
    document.add_paragraph("Расположите вещества в порядке возрастания массы.")
    _answer(document, None, "312")

    _task(document, 6)
    document.add_paragraph("Впишите название процесса.")
    _answer(document, "Переход из твёрдого состояния в газообразное.", "возгонка#сублимация")

    _task(document, 7)
    document.add_paragraph("Измерьте значение и запишите его с погрешностью.")
    _answer(document, None, "0,100,01")

    _task(document, 8)
    document.add_paragraph("Определите вещество по прибору.")
    document.add_picture(_png(), width=Inches(1))
    _answer(document, None, "42")

    document.save(str(path))


def build_repository(root: Path) -> Path:
    (root / "school" / "diagnostics").mkdir(parents=True)
    (root / "school" / "assets" / "questions").mkdir(parents=True)
    (root / "docs").mkdir()
    diagnostic = {
        "id": "oge-chemistry-1",
        "exam": "ОГЭ",
        "subject": "Химия",
        "mark": "9 класс",
        "quick_count": 1,
        "scoring": {"max_score": 100, "score_unit": "accuracy_percent"},
        "questions": [
            {
                "id": "seed1",
                "type": "input",
                "topic": "Периодический закон",
                "title": "Задание 4",
                "prompt": "Сколько протонов у углерода?",
                "max_primary_score": 1,
                "source": {
                    "provider": "maximum",
                    "official_year": 2026,
                    "approval_status": "draft",
                    "source_kind": "original",
                    "source_url": "https://maximumtest.ru/",
                    "exam_position": "4",
                    "rights_status": "original",
                    "verified_at": "2026-09-01",
                },
                "correct": ["6"],
            }
        ],
    }
    path = root / "school" / "diagnostics" / "oge-chemistry-1.json"
    path.write_text(
        json.dumps(diagnostic, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    (root / "docs" / "FIPI_2026_CONTENT_MATRIX.md").write_text(
        "## ОГЭ\n\n"
        "### Химия, `oge-chemistry-1`\n\n"
        "| Позиция КИМ | Проверяемое умение | Формат | Макс. балл | Надёжность |\n"
        "|---|---|---:|---:|---|\n"
        "| 4 | Периодический закон и строение атома | К, число | 2 | предварительно |\n",
        encoding="utf-8",
        newline="\n",
    )
    return path


@pytest.fixture()
def imported(tmp_path):
    source_directory = tmp_path / "docx"
    source_directory.mkdir()
    build_source_document(source_directory / SOURCE_NAME)
    catalog_path = build_repository(tmp_path)
    importer.main([str(source_directory), "--root", str(tmp_path)])
    return tmp_path, catalog_path, source_directory


def _questions(catalog_path: Path) -> dict[str, dict]:
    payload = json.loads(catalog_path.read_text(encoding="utf-8"))
    return {question["id"]: question for question in payload["questions"]}


def test_every_mapping_is_emitted_in_the_catalog_format(imported):
    _, catalog_path, _ = imported
    questions = _questions(catalog_path)

    single = questions["sp-chemistry-oge-2022-q1"]
    assert single["type"] == "single"
    assert [option["id"] for option in single["options"]] == ["a", "b", "c"]
    assert single["correct"] == "c"
    assert single["explanation"] == "Хлор стоит третьим в списке."

    multiple = questions["sp-chemistry-oge-2022-q2"]
    assert multiple["type"] == "multiple"
    assert multiple["selection_limit"] == 2
    assert multiple["correct"] == ["a", "c"]
    assert "explanation" not in multiple

    matching = questions["sp-chemistry-oge-2022-q3"]
    assert matching["type"] == "matching"
    assert [item["label"] for item in matching["items"]] == ["А) HCl", "Б) NaOH"]
    assert [option["id"] for option in matching["options"]] == ["o1", "o2"]
    assert matching["correct"] == {"i1": "o1", "i2": "o2"}
    assert "ФОРМУЛА" not in matching["prompt"]

    numeric = questions["sp-chemistry-oge-2022-q4"]
    assert numeric["type"] == "input"
    assert numeric["correct"] == ["0,25", "0.25"]

    sequence = questions["sp-chemistry-oge-2022-q5"]
    assert sequence["type"] == "input"
    assert sequence["correct"] == ["312"]
    assert sequence["prompt"].endswith(importer.SEQUENCE_HINT)

    text = questions["sp-chemistry-oge-2022-q6"]
    assert text["type"] == "text"
    assert text["correct"] == ["возгонка", "сублимация"]
    assert text["max_length"] == 80


def test_source_metadata_marks_editorial_drafts(imported):
    _, catalog_path, _ = imported
    source = _questions(catalog_path)["sp-chemistry-oge-2022-q1"]["source"]
    assert source == {
        "provider": "maximum_editorial",
        "official_year": 2022,
        "approval_status": "draft",
        "source_kind": "original",
        "source_url": "https://maximumtest.ru/",
        "rights_status": "original",
        "verified_at": source["verified_at"],
    }


def test_matrix_topic_applies_only_where_the_catalog_asserts_that_position(imported):
    _, catalog_path, _ = imported
    questions = _questions(catalog_path)
    mapped = questions["sp-chemistry-oge-2022-q4"]
    assert mapped["topic"] == "Периодический закон и строение атома"
    assert mapped["max_primary_score"] == 2
    assert mapped["source"]["exam_position"] == "4"

    unmapped = questions["sp-chemistry-oge-2022-q1"]
    assert unmapped["topic"] == "Задание 1"
    assert unmapped["max_primary_score"] == 1
    assert "exam_position" not in unmapped["source"]


def test_irregular_key_is_skipped_and_explained_in_the_report(imported):
    root, catalog_path, _ = imported
    assert "sp-chemistry-oge-2022-q7" not in _questions(catalog_path)
    report = (root / "authoring" / "sharepoint-import" / "report.md").read_text(
        encoding="utf-8"
    )
    assert SOURCE_NAME in report
    assert "| 7 | skipped | - | irregular_key | 0 | - |" in report
    assert "| 1 | imported | single | - | 0 | topic_unmapped |" in report
    assert "| 4 | imported | input | - | 0 | по матрице ФИПИ |" in report


def test_inline_figure_becomes_a_deduplicated_question_asset(imported):
    root, catalog_path, _ = imported
    question = _questions(catalog_path)["sp-chemistry-oge-2022-q8"]
    assert question["asset"] == "assets/questions/sp-chemistry-oge-2022-q8-1.png"
    assert "assets" not in question
    asset = root / "school" / question["asset"]
    with Image.open(asset) as image:
        assert image.size == (40, 30)


def test_existing_questions_keep_their_exact_bytes(imported):
    _, catalog_path, _ = imported
    text = catalog_path.read_text(encoding="utf-8")
    assert '"id": "seed1"' in text
    assert '"correct": [\n        "6"\n      ]' in text
    assert json.loads(text)["questions"][0]["id"] == "seed1"
    assert json.loads(text)["quick_count"] == 1


def test_rerunning_the_import_replaces_only_the_prefixed_questions(imported):
    root, catalog_path, source_directory = imported
    before = catalog_path.read_bytes()
    assets_before = {
        path.name: path.read_bytes()
        for path in (root / "school" / "assets" / "questions").iterdir()
    }

    importer.main([str(source_directory), "--root", str(root)])

    assert catalog_path.read_bytes() == before
    assert {
        path.name: path.read_bytes()
        for path in (root / "school" / "assets" / "questions").iterdir()
    } == assets_before


def test_stale_prefixed_assets_are_removed_on_reimport(imported):
    root, _, source_directory = imported
    stale = root / "school" / "assets" / "questions" / "sp-old-question-1.png"
    stale.write_bytes(_png(8, 8).getvalue())
    kept = root / "school" / "assets" / "questions" / "legacy.png"
    kept.write_bytes(_png(8, 8).getvalue())

    importer.main([str(source_directory), "--root", str(root)])

    assert not stale.exists()
    assert kept.exists()


def test_dry_run_leaves_the_repository_untouched(tmp_path):
    source_directory = tmp_path / "docx"
    source_directory.mkdir()
    build_source_document(source_directory / SOURCE_NAME)
    catalog_path = build_repository(tmp_path)
    before = catalog_path.read_bytes()

    importer.main([str(source_directory), "--root", str(tmp_path), "--dry-run"])

    assert catalog_path.read_bytes() == before
    assert not (tmp_path / "authoring").exists()


def test_repository_catalogs_round_trip_without_reformatting():
    for path in sorted((ROOT / "school" / "diagnostics").glob("*.json")):
        target = importer.read_target(path)
        assert importer.render_target(target, target.chunks) == path.read_text(
            encoding="utf-8"
        )


def test_unsupported_pdf_glyphs_are_normalized_instead_of_dropping_the_task():
    assert importer.clean_line("сто́л") == "стол"
    assert importer.clean_line("𝑚 · 𝑔") == "m · g"
    assert importer.clean_line("∠ABC") == "угол ABC"
    assert importer.renders(importer.clean_line("∠ABC = 30°"))
